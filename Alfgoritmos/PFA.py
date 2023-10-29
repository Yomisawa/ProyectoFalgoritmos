import argparse
import os
import tkinter as tk
from tkinter import messagebox
import mysql.connector
from docx import Document
import smtplib
from email.message import EmailMessage
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

entry_cliente_id = None
entry_producto_id = None
entry_cantidad = None
entry_venta_id = None  


# Establece la conexión a la base de datos MySQL
conexion = mysql.connector.connect(
    host="localhost",
    user="yeferson",
    password="987654321",
    port="3306",
    database="cliente_producto"
)

# Crea un objeto cursor para ejecutar consultas SQL
cursor = conexion.cursor()


# Función para calcular el total de ventas
def calcular_total_ventas():
    consulta = "SELECT SUM(cantidad) FROM ventas"
    cursor.execute(consulta)
    total_ventas = cursor.fetchone()[0]
    messagebox.showinfo("Total de Ventas", f"El total de ventas es: {total_ventas}")


# Funciones para clientes
def agregar_cliente():
    nombre = entry_nombre.get()
    direccion = entry_direccion.get()
    telefono = entry_telefono.get()
    
    consulta = "INSERT INTO clientes (id_cliente, nombre, direccion, telefono) VALUES (NULL, %s, %s, %s)"
    valores = (nombre, direccion, telefono)

    try:
        cursor.execute(consulta, valores)
        conexion.commit()
        cliente_id = cursor.lastrowid
        messagebox.showinfo("Éxito", "Cliente agregado exitosamente. ID del cliente: {}".format(cliente_id))
    except mysql.connector.Error as error:
        messagebox.showerror("Error", f"No se pudo agregar el cliente: {error}")
    
    clear_entries()

def ver_cliente():
    global cliente_id
    consulta = "SELECT id_cliente, nombre, direccion, telefono FROM clientes WHERE id_cliente = %s"
    valores = (entry_id_nuevo.get(),)

    cursor.execute(consulta, valores)
    cliente = cursor.fetchone()

    if cliente:
        mostrar_ventana_cliente(cliente)
    else:
        messagebox.showwarning("Cliente no encontrado", "No se encontró un cliente con el ID proporcionado.")
    
    clear_entries()

def mostrar_ventana_cliente(cliente):
    if cliente:
        ventana_cliente = tk.Toplevel()
        ventana_cliente.title("Información del Cliente")
        ventana_cliente.geometry("400x200")

        etiqueta_nombre = tk.Label(ventana_cliente, text=f"Nombre: {cliente[1]}", font=("Arial", 12))
        etiqueta_nombre.pack()

        etiqueta_direccion = tk.Label(ventana_cliente, text=f"Dirección: {cliente[2]}", font=("Arial", 12))
        etiqueta_direccion.pack()

        etiqueta_telefono = tk.Label(ventana_cliente, text=f"Teléfono: {cliente[3]}", font=("Arial", 12))
        etiqueta_telefono.pack()
    else:
        messagebox.showwarning("Cliente no encontrado", "No se encontró un cliente con el ID proporcionado.")

def eliminar_cliente():
    if entry_id_nuevo.get():  
        consulta = "DELETE FROM clientes WHERE id_cliente = %s"
        valores = (entry_id_nuevo.get(),)

        try:
            cursor.execute(consulta, valores)
            conexion.commit()
            messagebox.showinfo("Éxito", "Cliente eliminado exitosamente.")
        except mysql.connector.Error as error:
            messagebox.showerror("Error", f"No se pudo eliminar el cliente: {error}")
    else:
        messagebox.showwarning("ID no válido", "Por favor, proporcione un ID válido.")

    clear_entries()

def actualizar_cliente():
    if entry_id_nuevo.get():  # Verificar si se proporcionó un ID
        nombre = entry_nombre.get()
        direccion = entry_direccion.get()
        telefono = entry_telefono.get()

        consulta = "UPDATE clientes SET nombre = %s, direccion = %s, telefono = %s WHERE id_cliente = %s"
        valores = (nombre, direccion, telefono, entry_id_nuevo.get())

        try:
            cursor.execute(consulta, valores)
            conexion.commit()
            messagebox.showinfo("Éxito", "Cliente actualizado exitosamente.")
        except mysql.connector.Error as error:
            messagebox.showerror("Error", f"No se pudo actualizar el cliente: {error}")
    else:
        messagebox.showwarning("ID no válido", "Por favor, proporcione un ID válido.")

    clear_entries()


  # Funciones para productos
def agregar_producto():
    nombre = entry_nombre.get()
    precio = entry_precio.get()
    existencia = entry_existencia.get()
    proveedor = entry_proveedor.get()

    consulta = "INSERT INTO productos (nombre_producto, precio, existencia, proveedor) VALUES (%s, %s, %s, %s)"
    valores = (nombre, precio, existencia, proveedor)

    try:
        cursor.execute(consulta, valores)
        conexion.commit()
        id_producto = cursor.lastrowid  
        messagebox.showinfo("Éxito", "Producto agregado exitosamente. ID del producto: {}".format(id_producto))
    except mysql.connector.Error as error:
        messagebox.showerror("Error", f"No se pudo agregar el producto: {error}")
    
    clear_entries()

def ver_producto():
    global id_producto
    consulta = "SELECT id_producto, nombre_producto, precio, existencia, proveedor FROM productos WHERE id_producto = %s"
    valores = (entry_id_nuevo.get(),)

    cursor.execute(consulta, valores)
    producto = cursor.fetchone()

    if producto:
        mostrar_ventana_producto(producto)
    else:
        messagebox.showwarning("Producto no encontrado", "No se encontró un producto con el ID proporcionado.")
    
    clear_entries()

def mostrar_ventana_producto(producto):
    if producto:
        ventana_producto = tk.Toplevel()
        ventana_producto.title("Información del Producto")
        ventana_producto.geometry("300x100")

        etiqueta_nombre = tk.Label(ventana_producto, text=f"Nombre: {producto[1]}", font=("Arial", 12))
        etiqueta_nombre.pack()

        etiqueta_precio = tk.Label(ventana_producto, text=f"Precio: {producto[2]}", font=("Arial", 12))
        etiqueta_precio.pack()

        etiqueta_existencia= tk.Label(ventana_producto, text=f"existencia: {producto[3]}", font=("Arial", 12))
        etiqueta_existencia.pack()

        etiqueta_proveedor = tk.Label(ventana_producto, text=f"proveedor: {producto[4]}", font=("Arial", 12))
        etiqueta_proveedor.pack()
    else:
        messagebox.showwarning("Producto no encontrado", "No se encontró un producto con el ID proporcionado.")

def eliminar_producto():
    if entry_id_nuevo.get():  
        consulta = "DELETE FROM productos WHERE id_producto = %s"
        valores = (entry_id_nuevo.get(),)

        try:
            cursor.execute(consulta, valores)
            conexion.commit()
            messagebox.showinfo("Éxito", "Producto eliminado exitosamente.")
        except mysql.connector.Error as error:
            messagebox.showerror("Error", f"No se pudo eliminar el producto: {error}")
    else:
        messagebox.showwarning("ID no válido", "Por favor, proporcione un ID válido.")

    clear_entries()

def actualizar_producto():
    if entry_id_nuevo.get():  
        nombre = entry_nombre.get()
        precio = entry_precio.get()

        consulta = "UPDATE productos SET nombre_producto = %s, precio = %s WHERE id_producto = %s"
        valores = (nombre, precio, entry_id_nuevo.get())

        try:
            cursor.execute(consulta, valores)
            conexion.commit()
            messagebox.showinfo("Éxito", "Producto actualizado exitosamente.")
        except mysql.connector.Error as error:
            messagebox.showerror("Error", f"No se pudo actualizar el producto: {error}")
    else:
        messagebox.showwarning("ID no válido", "Por favor, proporcione un ID válido.")

 # Funciones para ventas
def agregar_venta():
    cliente_id = entry_cliente_id.get()
    producto_id = entry_producto_id.get()
    cantidad = entry_cantidad.get()

    if not cliente_id or not producto_id or not cantidad:
        messagebox.showwarning("Campos vacíos", "Por favor, complete todos los campos.")
        return

    # Verificar si el cliente y el producto existen
    cursor.execute("SELECT * FROM clientes WHERE id_cliente = %s", (cliente_id,))
    cliente = cursor.fetchone()
    cursor.execute("SELECT * FROM productos WHERE id_producto = %s", (producto_id,))
    producto = cursor.fetchone()

    try:
        cantidad = int(cantidad)  # Convertir la cantidad en un entero

        if cantidad <= 0:
            messagebox.showerror("Cantidad inválida", "La cantidad debe ser un valor positivo.")
            return

        if not cliente or not producto:
            messagebox.showerror("Error", "Cliente o producto no encontrado. Verifique los IDs.")
            return

        existencia = int(producto[3])  # Convertir la existencia en un entero

        # Verificar si hay suficiente existencia
        if cantidad > existencia:
            messagebox.showerror("Existencia insuficiente", "No hay suficiente existencia para este producto.")
            return
    except ValueError:
        messagebox.showerror("Valor inválido", "La cantidad no es un número válido.")
        return

    # Registrar la venta
    consulta = "INSERT INTO ventas (id_cliente, id_producto, cantidad) VALUES (%s, %s, %s)"
    valores = (cliente_id, producto_id, cantidad)

    try:
        cursor.execute(consulta, valores)
        conexion.commit()
        messagebox.showinfo("Venta registrada", "La venta se ha registrado exitosamente.")

        # Actualizar la existencia del producto
        nueva_existencia = existencia - cantidad
        cursor.execute("UPDATE productos SET existencia = %s WHERE id_producto = %s", (nueva_existencia, producto_id))
        conexion.commit()
    except mysql.connector.Error as error:
        messagebox.showerror("Error", f"No se pudo registrar la venta: {error}")

    clear_entries()

   # Función para anular una venta
def anular_venta():
    global entry_venta_id  
    venta_id_str = entry_venta_id.get()

    if not venta_id_str:
        messagebox.showwarning("Campo vacío", "Por favor, ingrese el ID de la venta a anular.")
        return

    try:
        venta_id = entry_venta_id.get()  # Obtener el valor como entero
        
    except ValueError:
        messagebox.showerror("Error", "El ID de la venta debe ser un número entero.")

    consulta = "DELETE FROM ventas WHERE id_venta = %s"
    valores = (venta_id,)

    try:
        cursor.execute(consulta, valores)
        conexion.commit()
        messagebox.showinfo("Venta anulada", f"Venta con ID {venta_id} anulada exitosamente.")
    except mysql.connector.Error as error:
        messagebox.showerror("Error", f"No se pudo anular la venta: {error}")

    clear_entries()

# Función para ver ventas
def ver_ventas():
    consulta = "SELECT * FROM ventas"
    cursor.execute(consulta)
    ventas = cursor.fetchall()
    
    if ventas:
        mostrar_ventas(ventas)
    else:
        messagebox.showinfo("Sin ventas", "No hay ventas registradas en la base de datos.")
    
    clear_entries()

# Función para ver ventas por cliente
def ventas_por_cliente():
    global entry_cliente_id
    cliente_id = entry_cliente_id.get()

    if not cliente_id:
        messagebox.showwarning("Campo vacío", "Por favor, ingrese el ID del cliente para ver sus ventas.")
        return

    consulta = "SELECT * FROM ventas WHERE id_cliente = %s"
    valores = (cliente_id,)

    cursor.execute(consulta, valores)
    ventas_cliente = cursor.fetchall()
    
    if ventas_cliente:
        mostrar_ventas(ventas_cliente)
    else:
        messagebox.showinfo("Sin ventas", "No hay ventas registradas para el cliente con ID {}".format(cliente_id))
    
    clear_entries()

# Función para ver ventas por producto
def ventas_por_producto():
    global entry_producto_id
    producto_id = entry_producto_id.get()

    if not producto_id:
        messagebox.showwarning("Campo vacío", "Por favor, ingrese el ID del producto para ver sus ventas.")
        return

    consulta = "SELECT * FROM ventas WHERE id_producto = %s"
    valores = (producto_id,)

    cursor.execute(consulta, valores)
    ventas_producto = cursor.fetchall()
    
    if ventas_producto:
        mostrar_ventas(ventas_producto)
    else:
        messagebox.showinfo("Sin ventas", "No hay ventas registradas para el producto con ID {}".format(producto_id))
    
    clear_entries()

# Función para mostrar las ventas en una ventana separada
def mostrar_ventas(ventas):
    ventana_ventas = tk.Toplevel()
    ventana_ventas.title("Lista de Ventas")
    ventana_ventas.geometry("500x300")
    
    lista_ventas = tk.Listbox(ventana_ventas, selectbackground="yellow", selectmode=tk.SINGLE, font=("Arial", 12))
    lista_ventas.pack(fill=tk.BOTH, expand=1)

    for venta in ventas:
        lista_ventas.insert(tk.END, f"ID Venta: {venta[0]}, ID Cliente: {venta[1]}, ID Producto: {venta[2]}, Cantidad: {venta[3]}")
 

def clear_entries():
    entry_id_nuevo.delete(0, "end")
    entry_nombre.delete(0, "end")
    entry_direccion.delete(0, "end")
    entry_telefono.delete(0, "end")
    entry_precio.delete(0, "end")
    entry_existencia.delete(0, "end")
    entry_proveedor.delete(0, "end")
    entry_cliente_id.delete(0, "end")
    entry_producto_id.delete(0, "end")
    entry_cantidad.delete(0, "end")

# Configuración de la interfaz gráfica
window = tk.Tk()
window.title("Sistema de Gestión")
window.geometry("800x300")
window.configure(bg="black")

def ventana_ventas():
    global entry_venta_id 
    ventana_ventas = tk.Toplevel()
    ventana_ventas.title("Ventas")
    ventana_ventas.geometry("300x200")
    
# Crear entradas para la venta
    global entry_cliente_id, entry_producto_id, entry_cantidad
    
    label_cliente_id = tk.Label(ventana_ventas, text="ID del Cliente:", font=("Arial", 12))
    label_cliente_id.pack()

    entry_cliente_id = tk.Entry(ventana_ventas, font=("Arial", 12))
    entry_cliente_id.pack()

    label_producto_id = tk.Label(ventana_ventas, text="ID del Producto:", font=("Arial", 12))
    label_producto_id.pack()

    entry_producto_id = tk.Entry(ventana_ventas, font=("Arial", 12))
    entry_producto_id.pack()

    label_cantidad = tk.Label(ventana_ventas, text="Cantidad:", font=("Arial", 12))
    label_cantidad.pack()

    entry_cantidad = tk.Entry(ventana_ventas, font=("Arial", 12))
    entry_cantidad.pack()

    btn_agregar_venta = tk.Button(ventana_ventas, text="Agregar Venta", command=agregar_venta, bg="white", fg="black", font=("Arial", 12))
    btn_agregar_venta.pack()
    
# Etiquetas y campos de entrada
label_id_nuevo = tk.Label(window, text="Cambios Por ID:", bg="black", fg="white", font=("Arial", 12))
label_id_nuevo.grid(row=0, column=2, padx=10, pady=10, sticky="w")

entry_id_nuevo = tk.Entry(window, font=("Arial", 12))
entry_id_nuevo.grid(row=0, column=3, padx=10, pady=10)

label_nombre = tk.Label(window, text="Nombre:", bg="black", fg="white", font=("Arial", 12))
label_nombre.grid(row=0, column=0, padx=10, pady=10, sticky="w")

entry_nombre = tk.Entry(window, font=("Arial", 12))
entry_nombre.grid(row=0, column=1, padx=10, pady=10)

label_direccion = tk.Label(window, text="Dirección:", bg="black", fg="white", font=("Arial", 12))
label_direccion.grid(row=1, column=0, padx=10, pady=10, sticky="w")

entry_direccion = tk.Entry(window, font=("Arial", 12))
entry_direccion.grid(row=1, column=1, padx=10, pady=10)

label_telefono = tk.Label(window, text="Teléfono:", bg="black", fg="white", font=("Arial", 12))
label_telefono.grid(row=2, column=0, padx=10, pady=10, sticky="w")

entry_telefono = tk.Entry(window, font=("Arial", 12))
entry_telefono.grid(row=2, column=1, padx=10, pady=10)

label_precio = tk.Label(window, text="Precio:", fg="white", bg="black", font=("Arial", 12))
label_precio.grid(row=3, column=0, padx=10, pady=10, sticky="w")

entry_precio = tk.Entry(window, font=("Arial", 12))
entry_precio.grid(row=3, column=1, padx=10, pady=10)

label_existencia = tk.Label(window, text="Existencia:", fg="white", bg="black", font=("Arial", 12))
label_existencia.grid(row=4, column=0, padx=10, pady=10, sticky="w")

entry_existencia = tk.Entry(window, font=("Arial", 12))
entry_existencia.grid(row=4, column=1, padx=10, pady=10)

label_proveedor = tk.Label(window, text="Proveedor:", fg="white", bg="black", font=("Arial", 12))
label_proveedor.grid(row=5, column=0, padx=10, pady=10, sticky="w")

entry_proveedor = tk.Entry(window, font=("Arial", 12))
entry_proveedor.grid(row=5, column=1, padx=10, pady=10)


def ventana_ventas():
    ventana_ventas = tk.Toplevel()
    ventana_ventas.title("Ventas")
    ventana_ventas.geometry("300x200")
    global entry_cliente_id, entry_producto_id, entry_cantidad, entry_venta_id 

# Crear entradas para la venta
    global entry_cliente_id, entry_producto_id, entry_cantidad
    
    label_cliente_id = tk.Label(ventana_ventas, text="ID del Cliente:", font=("Arial", 12))
    label_cliente_id.pack()

    entry_cliente_id = tk.Entry(ventana_ventas, font=("Arial", 12))
    entry_cliente_id.pack()

    label_producto_id = tk.Label(ventana_ventas, text="ID del Producto:", font=("Arial", 12))
    label_producto_id.pack()

    entry_producto_id = tk.Entry(ventana_ventas, font=("Arial", 12))
    entry_producto_id.pack()

    label_cantidad = tk.Label(ventana_ventas, text="Cantidad:", font=("Arial", 12))
    label_cantidad.pack()

    entry_cantidad = tk.Entry(ventana_ventas, font=("Arial", 12))
    entry_cantidad.pack()

    btn_agregar_venta = tk.Button(ventana_ventas, text="Agregar Venta", command=agregar_venta, bg="white", fg="black", font=("Arial", 12))
    btn_agregar_venta.pack()
    
    # Botones adicionales en la ventana de ventas
    label_venta_id = tk.Label(ventana_ventas, text="ID de Venta:", font=("Arial", 12))
    label_venta_id.pack()

    entry_venta_id = tk.Entry(ventana_ventas, font=("Arial", 12))
    entry_venta_id.pack()

    btn_anular_venta = tk.Button(ventana_ventas, text="Anular Venta", command=anular_venta, bg="red", fg="white", font=("Arial", 12))
    btn_anular_venta.pack()

    btn_ver_ventas = tk.Button(ventana_ventas, text="Ver Todas las Ventas", command=ver_ventas, bg="white", fg="black", font=("Arial", 12))
    btn_ver_ventas.pack()

    btn_ventas_por_cliente = tk.Button(ventana_ventas, text="Ventas por Cliente", command=ventas_por_cliente, bg="white", fg="black", font=("Arial", 12))
    btn_ventas_por_cliente.pack()

    btn_ventas_por_producto = tk.Button(ventana_ventas, text="Ventas por Producto", command=ventas_por_producto, bg="white", fg="black", font=("Arial", 12))
    btn_ventas_por_producto.pack()
    

# Botones
btn_agregar_cliente = tk.Button(window, text="Agregar Cliente", command=agregar_cliente, bg="white", fg="black", font=("Arial", 12))
btn_agregar_cliente.grid(row=9, column=0, padx=10, pady=10)

btn_ver_cliente = tk.Button(window, text="Ver Cliente", command=ver_cliente, bg="white", fg="black", font=("Arial", 12))
btn_ver_cliente.grid(row=9, column=1, padx=10, pady=10)

btn_eliminar_cliente = tk.Button(window, text="Eliminar Cliente", command=eliminar_cliente, bg="white", fg="black", font=("Arial", 12))
btn_eliminar_cliente.grid(row=9, column=2, padx=10, pady=10)

btn_actualizar_cliente = tk.Button(window, text="Actualizar Cliente", command=actualizar_cliente, bg="white", fg="black", font=("Arial", 12))
btn_actualizar_cliente.grid(row=9, column=3, padx=10, pady=10)

btn_agregar_producto = tk.Button(window, text="Agregar Producto", command=agregar_producto, bg="white", fg="black", font=("Arial", 12))
btn_agregar_producto.grid(row=10, column=0, padx=10, pady=10)

btn_ver_producto = tk.Button(window, text="Ver Producto", command=ver_producto, bg="white", fg="black", font=("Arial", 12))
btn_ver_producto.grid(row=10, column=1, padx=10, pady=10)

btn_eliminar_producto = tk.Button(window, text="Eliminar Producto", command=eliminar_producto, bg="white", fg="black", font=("Arial", 12))
btn_eliminar_producto.grid(row=10, column=2, padx=10, pady=10)

btn_actualizar_producto = tk.Button(window, text="Actualizar Producto", command=actualizar_producto, bg="white", fg="black", font=("Arial", 12))
btn_actualizar_producto.grid(row=10, column=3, padx=10, pady=10)

btn_ventas = tk.Button(window, text="Ventas", command=ventana_ventas, bg="white", fg="black", font=("Arial", 12))
btn_ventas.grid(row=8, column=2, padx=10, pady=10)


# Función para guardar registros en Word
def guardar_registros(tipo_registro):
    doc = Document()
    
    if tipo_registro == "clientes":
        consulta = "SELECT * FROM clientes"
        header = "ID Cliente - Nombre - Dirección - Teléfono"
    elif tipo_registro == "productos":
        consulta = "SELECT * FROM productos"
        header = "ID Producto - Nombre - Precio - Existencia - Proveedor"
    else:
        consulta = "SELECT * FROM ventas"
        header = "ID Venta - ID Cliente - ID Producto - Cantidad"
    
    cursor.execute(consulta)
    registros = cursor.fetchall()
    
    doc.add_heading(f"Registros de {tipo_registro}", 0)
    doc.add_paragraph(header)
    
    for registro in registros:
        registro_str = " - ".join(map(str, registro))
        doc.add_paragraph(registro_str)
    
    doc.save(f"registros_{tipo_registro}.docx")
    messagebox.showinfo("Registros Guardados", f"Los registros de {tipo_registro} se han guardado en registros_{tipo_registro}.docx")

# Función para enviar correos
def enviar_correo():
    # Crea una ventana emergente para ingresar detalles del correo
    ventana_correo = tk.Toplevel()
    ventana_correo.title("Enviar Correo")
    ventana_correo.geometry("400x300")

    # Agrega campos para el destinatario, el asunto y el cuerpo del correo
    label_destinatario = tk.Label(ventana_correo, text="Destinatario:")
    label_destinatario.pack()
    entry_destinatario = tk.Entry(ventana_correo, width=30)
    entry_destinatario.pack()

    label_asunto = tk.Label(ventana_correo, text="Asunto:")
    label_asunto.pack()
    entry_asunto = tk.Entry(ventana_correo, width=30)
    entry_asunto.pack()

    label_cuerpo = tk.Label(ventana_correo, text="Cuerpo:")
    label_cuerpo.pack()
    entry_cuerpo = tk.Text(ventana_correo, width=30, height=10)
    entry_cuerpo.pack()

    # Lista de archivos DOCX en el directorio actual
    archivos_docx = [f for f in os.listdir('.') if f.endswith('.docx')]
    
    # Campo desplegable para seleccionar el archivo a adjuntar
    label_adjunto = tk.Label(ventana_correo, text="Archivo adjunto:")
    label_adjunto.pack()
    
    archivo_adjunto = tk.StringVar(ventana_correo)
    archivo_adjunto.set(archivos_docx[0])  # valor por defecto
    
    menu_adjunto = tk.OptionMenu(ventana_correo, archivo_adjunto, *archivos_docx)
    menu_adjunto.pack()

    # Función para enviar el correo
    def enviar():
        destinatario = entry_destinatario.get()
        asunto = entry_asunto.get()

        # Crear un objeto EmailMessage
        mensaje = EmailMessage()
        mensaje.set_content(entry_cuerpo.get("1.0", "end"))

        mensaje["Subject"] = asunto
        mensaje["From"] = "tecgt80@gmail.com"
        mensaje["To"] = destinatario

        # Adjuntar el archivo DOCX seleccionado
        with open(archivo_adjunto.get(), "rb") as f:
            mensaje.add_attachment(f.read(), maintype="application", subtype="vnd.openxmlformats-officedocument.wordprocessingml.document", filename=archivo_adjunto.get())

        # Conectar al servidor SMTP de Google (debes permitir aplicaciones menos seguras en tu cuenta de Google)
        try:
            server = smtplib.SMTP("smtp.gmail.com", 587)
            server.starttls()
            server.login("tecgt80@gmail.com", "********")

            server.send_message(mensaje)
            server.quit()

            messagebox.showinfo("Correo enviado", "El correo se ha enviado con éxito.")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo enviar el correo: {e}")

        ventana_correo.destroy()

    # Agregar un botón para enviar el correo
    btn_enviar = tk.Button(ventana_correo, text="Enviar", command=enviar, bg="white", fg="black", font=("Arial", 12))
    btn_enviar.pack()

 # Botones para las nuevas funciones
btn_total_ventas = tk.Button(window, text="Total de Ventas", command=calcular_total_ventas, bg="white", fg="black", font=("Arial", 12))
btn_total_ventas.grid(row=8, column=3, padx=10, pady=10)

btn_guardar_clientes = tk.Button(window, text="Guardar Registros de Clientes", command=lambda: guardar_registros("clientes"), bg="white", fg="black", font=("Arial", 12))
btn_guardar_clientes.grid(row=7, column=0, padx=10, pady=10)

btn_guardar_productos = tk.Button(window, text="Guardar Registros de Productos", command=lambda: guardar_registros("productos"), bg="white", fg="black", font=("Arial", 12))
btn_guardar_productos.grid(row=7, column=1, padx=10, pady=10)

btn_guardar_ventas = tk.Button(window, text="Guardar Registros de Ventas", command=lambda: guardar_registros("ventas"), bg="white", fg="black", font=("Arial", 12))
btn_guardar_ventas.grid(row=7, column=2, padx=10, pady=10)

btn_enviar_correo = tk.Button(window, text="Enviar Correo", command=enviar_correo, bg="white", fg="black", font=("Arial", 12))
btn_enviar_correo.grid(row=8, column=3, padx=10, pady=10)
btn_enviar_correo = tk.Button(window, text="Enviar Correo", command=enviar_correo, bg="white", fg="black", font=("Arial", 12))
btn_enviar_correo.grid(row=8, column=3, padx=10, pady=10)

window.mainloop()
# Cerrar el cursor y la conexión al final del programa
cursor.close()
conexion.close()
